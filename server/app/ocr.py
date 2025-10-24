# # server/app/ocr.py
# import fitz, hashlib
# from docx import Document
# # Optional: import easyocr

# def make_doc_id(name: str, content: bytes) -> str:
#     return hashlib.sha1((name+str(len(content))).encode()).hexdigest()[:12]

# def read_pdf(bytestr: bytes) -> str:
#     doc = fitz.open(stream=bytestr, filetype="pdf")
#     return "\n\n".join([p.get_text() for p in doc])

# def read_docx(bytestr) -> str:
#     d = Document(bytestr)
#     return "\n".join(p.text for p in d.paragraphs)

# def read_any(name: str, content: bytes) -> str:
#     if name.lower().endswith(".pdf"):
#         txt = read_pdf(content)
#         # if too sparse → fallback to OCR per page with EasyOCR (optional for time)
#         return txt
#     if name.lower().endswith(".docx"):
#         return read_docx(content)
#     raise ValueError("Unsupported file type")


# ──────────────────────────────────────────────────────────────────────────────
# Repo: server/app
# Files in this single canvas:
#   - main.py
#   - models.py
#   - ocr.py
#   - prompts.py
#   - extract.py
#   - requirements.txt (at server/)
#   - run.sh
# Copy each section into its respective file path.
# ──────────────────────────────────────────────────────────────────────────────

# ============================= ocr.py ========================================
import io
import re
import hashlib
from typing import Tuple, Dict

# PDF text extractor
import fitz  # PyMuPDF

# DOCX
from docx import Document

# Optional OCR fallback (disable if not installed or torch not available)
_HAS_EASYOCR = False
try:
    import easyocr  # heavier; supports en/ar
    try:
        # try to create a lightweight reader quickly to confirm torch works
        _ = easyocr.Reader(['en'], gpu=False)
        _HAS_EASYOCR = True
    except Exception:
        _HAS_EASYOCR = False
except Exception:
    _HAS_EASYOCR = False


def make_doc_id(name: str, content: bytes) -> str:
    return hashlib.sha1((name + str(len(content))).encode()).hexdigest()[:12]


def _pdf_text(bytestr: bytes) -> Tuple[str, Dict]:
    doc = fitz.open(stream=bytestr, filetype="pdf")
    pages = []
    empty_pages = 0
    for i in range(len(doc)):
        t = doc[i].get_text()
        if len(t.strip()) < 10:
            empty_pages += 1
        pages.append(t)
    text = "\n\f\n".join(pages)  # keep page breaks
    return text, {"pages": len(doc), "empty_pages": empty_pages}


def _pdf_ocr(bytestr: bytes, langs=("en", "ar")) -> Tuple[str, Dict]:
    if not _HAS_EASYOCR:
        return "", {"ocr": False}
    reader = easyocr.Reader(list(langs), gpu=False)
    doc = fitz.open(stream=bytestr, filetype="pdf")
    pages_text = []
    for i in range(len(doc)):
        pix = doc[i].get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        res = reader.readtext(img_bytes, detail=0, paragraph=True)
        pages_text.append("\n".join(res))
    return "\n\f\n".join(pages_text), {"ocr": True, "pages": len(doc)}


def _docx_text(bytestr: bytes) -> Tuple[str, Dict]:
    # python-docx accepts a file-like
    bio = io.BytesIO(bytestr)
    d = Document(bio)
    text = "\n".join(p.text for p in d.paragraphs)
    return text, {"pages": None}


def _text_density(s: str) -> float:
    if not s:
        return 0.0
    letters = sum(ch.isalnum() for ch in s)
    return letters / max(1, len(s))


def _has_low_text(text: str, meta: Dict) -> bool:
    # Heuristic: few characters or many empty pages => try OCR
    return (len(text.strip()) < 200) or (meta.get("empty_pages", 0) >= max(1, meta.get("pages", 1)//2))


def read_any(name: str, content: bytes) -> Tuple[str, Dict]:
    name_l = name.lower()
    if name_l.endswith(".pdf"):
        text, meta = _pdf_text(content)
        if _has_low_text(text, meta):
            ocr_text, ocr_meta = _pdf_ocr(content)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
                meta.update(ocr_meta)
        # normalize whitespace
        text = re.sub(r"[ \t\u200f\u200e]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text, meta
    elif name_l.endswith(".docx"):
        text, meta = _docx_text(content)
        text = re.sub(r"[ \t\u200f\u200e]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text, meta
    else:
        raise ValueError("Unsupported file type (use PDF or DOCX)")


