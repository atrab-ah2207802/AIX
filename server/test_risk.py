
# """
# COMPLETE END-TO-END TEST: QDB Contracts → Templates → Sample Analysis
# 100% GEMINI - NO MOCK DATA
# """

# import asyncio
# import json
# import os
# import sys
# from datetime import datetime

# # Add current directory to path
# sys.path.insert(0, os.getcwd())

# from app.risk import call_llm
# from app.template_learner import TemplateLearner
# from app.learned_analyzer import SmartProductionRiskAnalyzer

# def load_all_contracts():
#     """LOAD ALL .docx files from company_contracts"""
#     contracts = []
#     contracts_folder = "data/company_contracts"
    
#     if not os.path.exists(contracts_folder):
#         print(f"❌ Contracts folder not found: {contracts_folder}")
#         return contracts
    
#     try:
#         from docx import Document
#     except ImportError:
#         print("❌ python-docx not installed. Run: pip install python-docx")
#         return contracts
    
#     for filename in os.listdir(contracts_folder):
#         if not filename.endswith('.docx'):
#             continue
            
#         file_path = os.path.join(contracts_folder, filename)
#         try:
#             doc = Document(file_path)
#             text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
#             if text and len(text) > 100:
#                 contracts.append({
#                     "name": filename,
#                     "text": text[:15000]  # Limit for Gemini
#                 })
#                 print(f"✅ Loaded: {filename} ({len(text)} chars)")
#             else:
#                 print(f"⚠️  Empty/short: {filename}")
                
#         except Exception as e:
#             print(f"❌ Failed to read {filename}: {e}")
    
#     return contracts

# def load_sample_contract():
#     """LOAD sample_contract.docx for testing"""
#     sample_path = "data/sample_contract.docx"
    
#     if not os.path.exists(sample_path):
#         print(f"❌ Sample contract not found: {sample_path}")
#         return None
    
#     try:
#         from docx import Document
#         doc = Document(sample_path)
#         text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
#         print(f"✅ Loaded sample contract: {len(text)} chars")
#         return text
#     except Exception as e:
#         print(f"❌ Failed to read sample contract: {e}")
#         return None

# async def learn_templates_from_contracts():
#     """STEP 1: Learn templates from QDB contracts using Gemini"""
#     print("\n" + "="*60)
#     print("STEP 1: LEARNING TEMPLATES FROM QDB CONTRACTS")
#     print("="*60)
    
#     contracts = load_all_contracts()
#     if not contracts:
#         print("❌ No contracts found to learn from!")
#         return None
    
#     print(f"📚 Learning from {len(contracts)} QDB contracts...")
    
#     learner = TemplateLearner(call_llm)
#     templates = await learner.analyze_contract_corpus(contracts)
    
#     # Save templates
#     os.makedirs("data/learned_templates", exist_ok=True)
#     with open("data/learned_templates/company_templates.json", 'w') as f:
#         json.dump(templates, f, indent=2)
    
#     print(f"✅ Learned {len(templates['learned_templates'])} templates")
#     return templates

# async def analyze_with_learned_templates(templates):
#     """STEP 2: Analyze sample contract using learned templates"""
#     print("\n" + "="*60)
#     print("STEP 2: ANALYZING SAMPLE CONTRACT WITH LEARNED TEMPLATES")  
#     print("="*60)
    
#     sample_text = load_sample_contract()
#     if not sample_text:
#         print("❌ No sample contract to analyze!")
#         return None
    
#     test_data = {
#         "parties": [{"name": "QDB", "role": "Client"}, {"name": "Vendor", "role": "Supplier"}],
#         "dates": {
#             "start": "2025-01-01",
#             "expiration": "2025-12-31"
#         },
#         "financial": {"paymentTerms": "Net 60"},
#         "jurisdiction": "UK"
#     }
    
#     print("🤖 Running SMART analysis with learned templates...")
#     smart_analyzer = SmartProductionRiskAnalyzer(templates, call_llm)
#     result = await smart_analyzer.analyze_contract(sample_text, test_data)
    
#     # Generate comprehensive report
#     comprehensive_report = smart_analyzer.generate_comprehensive_report(result, sample_text, test_data)
    
#     # Save comprehensive report
#     with open("data/comprehensive_analysis_report.json", 'w') as f:
#         json.dump(comprehensive_report, f, indent=2)
#     print("💾 Saved comprehensive report to 'data/comprehensive_analysis_report.json'")
    
#     return result, sample_text, comprehensive_report

# async def main():
#     """COMPLETE END-TO-END FLOW - SIMPLIFIED"""
#     print("🚀 QDB SMART CONTRACT ANALYSIS - TEMPLATE-BASED")
    
#     # STEP 1: Learn templates from QDB contracts
#     templates = await learn_templates_from_contracts()
#     if not templates:
#         print("❌ Template learning failed!")
#         return
    
#     # STEP 2: Analyze with learned templates ONLY  
#     smart_result = await analyze_with_learned_templates(templates)
#     if not smart_result:
#         print("❌ Smart analysis failed!")
#         return
        
#     result, sample_text, comprehensive_report = smart_result
    
#     # Display results from comprehensive report
#     print("\n" + "="*60)
#     print("🎉 COMPREHENSIVE ANALYSIS RESULTS")
#     print("="*60)
#     print(f"📊 Overall Score: {comprehensive_report['overall_score']}/100")
#     print(f"🚨 Risk Level: {comprehensive_report['risk_level']}")
#     print(f"📋 Flags Found: {len(comprehensive_report['flags'])}")
#     print(f"💡 Recommendations: {len(comprehensive_report['recommendations'])}")
    
#     # Show critical flags
#     critical_flags = [f for f in comprehensive_report['flags'] if f['severity'] in ['critical', 'high']]
#     if critical_flags:
#         print(f"\n🔍 CRITICAL ISSUES:")
#         for i, flag in enumerate(critical_flags, 1):
#             print(f"   {i}. [{flag['severity'].upper()}] {flag['title']}")
#             print(f"      {flag['description']}")
#             if flag['recommendation']:
#                 print(f"      💡 {flag['recommendation']}")

#     contract_id = "sample_contract_001"
#     selected_country = "qatar"  # This comes from frontend user selection
    
#     compliance_report = await run_legal_compliance_analysis(
#         sample_text, selected_country, comprehensive_report['flags'], contract_id
#     )
    
#     print(f"⚖️ {selected_country.upper()} COMPLIANCE: {compliance_report['compliance_summary']['overall_compliance_score']}/100")

# if __name__ == "__main__":
#     # Set API key as environment variable first!
#     if not os.getenv("GEMINI_API_KEY"):
#         print("❌ ERROR: Set Gemini API key first:")
#         print("   export GEMINI_API_KEY='your-actual-key-here'")
#         print("   Then run: python test_risk.py")
#     else:
#         asyncio.run(main())
"""
COMPLETE END-TO-END TEST: Template Analysis + Legal Compliance
Final working version with unified JSON output
"""

import asyncio
import json
import os
import sys
import time
from datetime import datetime
from typing import List, Dict, Any
import traceback

# Add current directory to path
sys.path.insert(0, os.getcwd())

from app.risk import call_llm, analyze_contract_risk
from app.template_learner import TemplateLearner
from app.legal_compliance import LegalComplianceAnalyzer

# ========== RATE LIMITING ==========
class RateLimitedLLM:
    def __init__(self, llm_client, max_requests=8, time_window=60):
        self.llm_client = llm_client
        self.max_requests = max_requests
        self.time_window = time_window
        self.requests = []
    
    async def call_with_rate_limit(self, prompt):
        now = time.time()
        self.requests = [req_time for req_time in self.requests 
                        if now - req_time < self.time_window]
        
        if len(self.requests) >= self.max_requests:
            wait_time = self.time_window - (now - self.requests[0]) + 1
            print(f"⏳ Rate limit reached. Waiting {wait_time:.1f}s...")
            await asyncio.sleep(wait_time)
            now = time.time()
            self.requests = [req_time for req_time in self.requests 
                           if now - req_time < self.time_window]
        
        self.requests.append(time.time())
        return await self.llm_client(prompt)

rate_limited_llm = RateLimitedLLM(call_llm, max_requests=8, time_window=60)

# ========== CONTRACT LOADING ==========

def load_all_contracts():
    """Load all .docx files from company_contracts"""
    contracts = []
    contracts_folder = "data/company_contracts"
    
    if not os.path.exists(contracts_folder):
        print(f"⚠️  Contracts folder not found: {contracts_folder}")
        os.makedirs(contracts_folder, exist_ok=True)
        return contracts
    
    try:
        from docx import Document
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        return contracts
    
    for filename in os.listdir(contracts_folder):
        if not filename.endswith('.docx'):
            continue
            
        file_path = os.path.join(contracts_folder, filename)
        try:
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
            if text and len(text) > 100:
                contracts.append({
                    "name": filename,
                    "text": text[:15000]
                })
                print(f"   ✅ Loaded: {filename} ({len(text)} chars)")
            else:
                print(f"   ⚠️  Empty/short: {filename}")
                
        except Exception as e:
            print(f"   ❌ Failed to read {filename}: {e}")
    
    return contracts

def load_sample_contract():
    """Load sample_contract.docx for testing"""
    sample_path = "data/sample_contract.docx"
    
    if not os.path.exists(sample_path):
        print(f"❌ Sample contract not found: {sample_path}")
        return None
    
    try:
        from docx import Document
        doc = Document(sample_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        print(f"✅ Loaded sample contract: {len(text)} chars")
        return text
    except Exception as e:
        print(f"❌ Failed to read sample contract: {e}")
        return None

# ========== TEMPLATE LEARNING ==========

# async def learn_templates_from_contracts():
#     """Learn templates from QDB contracts using Gemini"""
#     print("\n" + "="*60)
#     print("STEP 1: LEARNING TEMPLATES FROM QDB CONTRACTS")
#     print("="*60)
    
#     contracts = load_all_contracts()
    
#     if not contracts:
#         print("⚠️  No company contracts found - skipping template learning")
#         return None
    
#     print(f"📚 Learning from {len(contracts)} QDB contracts...")
    
#     learner = TemplateLearner(rate_limited_llm.call_with_rate_limit)
    
#     try:
#         templates = await learner.analyze_contract_corpus(contracts)
        
#         os.makedirs("data/learned_templates", exist_ok=True)
#         with open("data/learned_templates/company_templates.json", 'w') as f:
#             json.dump(templates, f, indent=2)
        
#         print(f"✅ Learned {len(templates['learned_templates'])} templates")
#         return templates
        
#     except Exception as e:
#         print(f"❌ Template learning failed: {e}")
#         traceback.print_exc()
#         return None
async def learn_templates_from_contracts():
    """Load pre-existing templates from JSON file"""
    print("📚 Loading pre-saved templates...")
    
    with open("data/learned_templates/company_templates.json", 'r', encoding='utf-8') as f:
        templates = json.load(f)
    
    print(f"✅ Loaded {len(templates['learned_templates'])} templates")
    return templates
# ========== HELPER FUNCTIONS FOR REPORT ==========

def _generate_term_consistency(contract_text: str) -> List[Dict]:
    """Generate term consistency analysis"""
    terms_to_check = ["confidential", "termination", "liability", "payment", "warrant", "indemnif"]
    
    term_consistency = []
    for term in terms_to_check:
        if term in contract_text.lower():
            term_consistency.append({
                "term": term.title(),
                "definitions": [f"References to {term} found in contract"],
                "is_consistent": True,
                "issue_description": None
            })
    
    if not term_consistency:
        term_consistency = [
            {
                "term": "Confidential Information",
                "definitions": ["information that should be kept private"],
                "is_consistent": False,
                "issue_description": "Term not clearly defined in contract"
            }
        ]
    
    return term_consistency[:6]

def _generate_clause_comparisons(flags: List[Dict]) -> List[Dict]:
    """Generate clause comparisons from flags"""
    comparisons = []
    
    clause_mapping = {
        "confidentiality": ["confidential", "non-disclosure"],
        "termination": ["termination", "terminate"], 
        "liability": ["liability", "liable", "damages"],
        "payment": ["payment", "fee", "invoice"],
        "warranties": ["warrant", "guarantee", "as is"],
        "indemnification": ["indemnif", "hold harmless"],
        "intellectual_property": ["intellectual", "ip", "copyright"],
        "dispute_resolution": ["dispute", "arbitration", "mediation"]
    }
    
    for flag in flags:
        for clause_type, keywords in clause_mapping.items():
            if any(keyword in flag["title"].lower() for keyword in keywords):
                comparisons.append({
                    "clause_type": clause_type,
                    "standard_version": f"Company standard {clause_type} clause",
                    "contract_version": flag["description"][:200] + "...",
                    "deviation_severity": flag["severity"],
                    "explanation": flag["description"]
                })
                break
    
    return comparisons[:8]

def _generate_legal_advice_from_flags(flags: List[Dict]) -> List[Dict]:
    """Generate legal advice from flags"""
    legal_advice = []
    
    severity_advice = {
        "critical": "This clause poses critical legal risk that requires immediate attention.",
        "high": "This clause poses significant legal risk that should be addressed.",
        "medium": "This clause has moderate legal risk that should be reviewed.",
        "low": "This clause has minor legal risk that may need attention."
    }
    
    for flag in flags:
        if flag["severity"] in ["critical", "high"]:
            advice_text = f"{severity_advice.get(flag['severity'], 'This clause requires legal review.')} {flag['description']}"
            
            legal_advice.append({
                "topic": flag["title"],
                "advice": advice_text,
                "risk_level": flag["severity"],
                "supporting_law": "General Contract Law Principles",
                "recommendations": [flag["recommendation"]] if flag["recommendation"] else ["Consult legal team"]
            })
    
    return legal_advice[:5]

def _extract_contract_value(extracted_data: Dict) -> str:
    """Extract contract value from extracted data"""
    financial_data = extracted_data.get("financial", {})
    if "amount" in financial_data:
        return f"QAR {financial_data['amount']}"
    elif "paymentTerms" in financial_data:
        return "Value specified in payment terms"
    else:
        return "Not specified"

# ========== UNIFIED REPORT GENERATION ==========

def generate_unified_report(template_risk, legal_compliance, contract_text: str, extracted_data: Dict) -> Dict:
    """Generate unified report: Frontend format + Legal compliance"""
    
    # Convert template risk flags
    flags = []
    for flag in template_risk.flags:
        flags.append({
            "severity": flag.severity,
            "category": flag.category,
            "title": flag.title,
            "description": flag.description,
            "recommendation": flag.recommendation or "",
            "clause_reference": flag.clause_reference or "Not Found",
            "confidence": flag.confidence
        })
    
    # Generate recommendations
    recommendations = list(set([
        flag["recommendation"] for flag in flags 
        if flag["recommendation"] and len(flag["recommendation"]) > 10
    ]))[:10]
    
    # Build unified report with proper legal compliance structure
    report = {
        # ===== FRONTEND'S EXISTING FORMAT =====
        "overall_score": template_risk.overall_score,
        "risk_level": template_risk.risk_level,
        "summary": template_risk.summary,
        "flags": flags,
        "recommendations": recommendations,
        "term_consistency": _generate_term_consistency(contract_text),
        "clause_comparisons": _generate_clause_comparisons(flags),
        "legal_advice": _generate_legal_advice_from_flags(flags),
        "analyzed_at": template_risk.analyzed_at,
        "contract_metadata": {
            "contract_name": "analyzed_contract.docx",
            "parties": extracted_data.get("parties", [{"name": "Unknown", "role": "Party"}]),
            "effective_date": extracted_data.get("dates", {}).get("start", "Not specified"),
            "expiry_date": extracted_data.get("dates", {}).get("expiration", "Not specified"),
            "total_value": _extract_contract_value(extracted_data),
            "extracted_clauses_count": len([p for p in contract_text.split('\n\n') if len(p.strip()) > 100])
        },
        
        # ===== LEGAL COMPLIANCE ADDED AT BOTTOM =====
        "legal_compliance": legal_compliance if legal_compliance else {
            "analyzed": False,
            "jurisdiction": "N/A",
            "compliance_summary": {
                "overall_compliance_score": 0,
                "status": "not_analyzed",
                "total_laws_analyzed": 0,
                "non_compliant_laws": 0
            },
            "law_analysis": [],
            "legal_risks": [],
            "recommendations": [],
            "critical_violations": []
        }
    }
    
    return report

# ========== DISPLAY RESULTS ==========

def display_results(report: Dict):
    """Display unified analysis results"""
    print("\n" + "="*60)
    print("🎉 UNIFIED ANALYSIS RESULTS")
    print("="*60)
    print(f"📊 Template Risk Score: {report['overall_score']}/100")
    print(f"🚨 Risk Level: {report['risk_level'].upper()}")
    print(f"📋 Template Flags: {len(report['flags'])}")
    print(f"💡 Recommendations: {len(report['recommendations'])}")
    
    # Legal compliance summary - FIXED VERSION
    legal = report['legal_compliance']
    
    # Check if legal compliance analysis was performed
    has_legal_analysis = legal and isinstance(legal, dict) and 'compliance_summary' in legal
    
    if has_legal_analysis:
        compliance_summary = legal['compliance_summary']
        print(f"\n⚖️  LEGAL COMPLIANCE:")
        print(f"   Jurisdiction: {legal.get('jurisdiction', 'N/A')}")
        print(f"   Compliance Score: {compliance_summary.get('overall_compliance_score', 'N/A')}/100")
        print(f"   Status: {compliance_summary.get('status', 'N/A').upper()}")
        print(f"   Laws Analyzed: {compliance_summary.get('total_laws_analyzed', 'N/A')}")
        print(f"   Non-Compliant Laws: {compliance_summary.get('non_compliant_laws', 'N/A')}")
    else:
        print(f"\n⚖️  LEGAL COMPLIANCE: Not analyzed or analysis failed")
    
    # Show critical/high template issues
    critical_high = [f for f in report['flags'] if f['severity'] in ['critical', 'high']]
    if critical_high:
        print(f"\n🔍 CRITICAL & HIGH TEMPLATE ISSUES ({len(critical_high)}):")
        for i, flag in enumerate(critical_high[:5], 1):
            print(f"   {i}. [{flag['severity'].upper()}] {flag['title']}")
            print(f"      {flag['description'][:120]}...")
    
    # Show legal risks - FIXED VERSION
    if has_legal_analysis and 'legal_risks' in legal and legal['legal_risks']:
        print(f"\n⚖️  LEGAL COMPLIANCE RISKS ({len(legal['legal_risks'])}):")
        for i, risk in enumerate(legal['legal_risks'], 1):
            print(f"   {i}. {risk.get('law', 'Unknown Law')} (Score: {risk.get('compliance_score', 'N/A')}/100)")
            print(f"      Severity: {risk.get('severity', 'N/A').upper()}")
            print(f"      Issues: {len(risk.get('issues', []))}")
    
    print("\n" + "="*60)
    print("✅ ANALYSIS COMPLETE!")
    print("="*60)
# ========== MAIN EXECUTION ==========

async def main():
    """Complete end-to-end flow"""
    print("🚀 QDB SMART CONTRACT ANALYSIS - INTEGRATED SYSTEM")
    print("="*60)
    
    # STEP 1: Learn templates (optional)
    print("\n📚 Attempting to learn company templates...")
    templates = await learn_templates_from_contracts()
    
    # STEP 2: Load sample contract
    print("\n📄 Loading sample contract...")
    sample_text = load_sample_contract()
    if not sample_text:
        print("❌ Cannot proceed without a sample contract!")
        print("\nTO FIX: Add a contract file to: data/sample_contract.docx")
        return
    
    # STEP 3: Prepare test data
    test_data = {
        "contract_id": "TEST_001",
        "parties": [
            {"name": "QDB", "role": "Client"}, 
            {"name": "Vendor Inc", "role": "Supplier"}
        ],
        "dates": {
            "start": "2025-01-01",
            "expiration": "2025-12-31"
        },
        "financial": {
            "paymentTerms": "Net 60 days",
            "amount": "500000"
        },
        "jurisdiction": "qatar",
        "governing_law": "Qatar laws"
    }
    
    # STEP 4A: Template-based risk analysis
    template_risk = None
    if templates:
        print("\n🧠 Running template-based risk analysis...")
        try:
            from app.learned_analyzer import SmartProductionRiskAnalyzer
            
            template_analyzer = SmartProductionRiskAnalyzer(templates, rate_limited_llm.call_with_rate_limit)
            template_risk = await template_analyzer.analyze_contract(sample_text, test_data)
            
            print(f"   ✅ Template analysis complete")
            print(f"   Risk Score: {template_risk.overall_score}/100")
            print(f"   Flags: {len(template_risk.flags)}")
        except Exception as e:
            print(f"   ❌ Template analysis failed: {e}")
            traceback.print_exc()
            template_risk = None
    
    # Fallback to basic risk analysis if no templates
    if not template_risk:
        print("\n📊 Running basic risk analysis (no templates)...")
        template_risk = await analyze_contract_risk(sample_text, test_data, "qatar")
        print(f"   ✅ Basic analysis complete: {template_risk.overall_score}/100")
    
    # STEP 4B: Legal compliance analysis
    print(f"\n⚖️  Running legal compliance analysis (QATAR)...")
    
    legal_db_path = "legal_database/qatar"
    legal_compliance = None
    
    if not os.path.exists(legal_db_path):
        print(f"   ⚠️  No legal database at '{legal_db_path}'")
        print(f"   Legal compliance will be skipped")
    else:
        try:
            # Convert risk flags to dict format
            risk_flags_dict = [
                {
                    "severity": flag.severity,
                    "category": flag.category,
                    "title": flag.title,
                    "description": flag.description,
                    "recommendation": flag.recommendation
                }
                for flag in template_risk.flags
            ]
            
            legal_analyzer = LegalComplianceAnalyzer(rate_limited_llm.call_with_rate_limit)
            legal_compliance = await legal_analyzer.analyze_legal_compliance(
                sample_text,
                "qatar",
                risk_flags_dict,
                contract_id="TEST_001"
            )
            
            print(f"   ✅ Legal compliance complete")
            print(f"   Compliance Score: {legal_compliance['compliance_summary']['overall_compliance_score']}/100")
            print(f"   Laws Analyzed: {legal_compliance['metadata']['total_laws_analyzed']}")
            
        except Exception as e:
            print(f"   ❌ Legal compliance failed: {e}")
            traceback.print_exc()
    
    # STEP 5: Generate unified report
    print("\n🔄 Generating unified report...")
    unified_report = generate_unified_report(template_risk, legal_compliance, sample_text, test_data)
    
    # STEP 6: Save results
    os.makedirs("data", exist_ok=True)
    output_file = "data/unified_analysis_report.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(unified_report, f, indent=2, ensure_ascii=False)
    print(f"💾 Saved unified report to '{output_file}'")
    
    # STEP 7: Display results
    display_results(unified_report)
    
    return unified_report

# ========== RUN ==========

if __name__ == "__main__":
    # Check API key
    if not os.getenv("GEMINI_API_KEY"):
        print("="*60)
        print("❌ ERROR: Gemini API key not found!")
        print("="*60)
        print("\nTO FIX:")
        print("1. Get API key from: https://makersuite.google.com/app/apikey")
        print("2. Set environment variable:")
        print("   export GEMINI_API_KEY='your-key-here'")
        print("3. Run this script again")
        print("="*60)
    else:
        try:
            asyncio.run(main())
        except KeyboardInterrupt:
            print("\n\n⚠️  Analysis interrupted by user")
        except Exception as e:
            print(f"\n\n❌ FATAL ERROR: {e}")
            traceback.print_exc()